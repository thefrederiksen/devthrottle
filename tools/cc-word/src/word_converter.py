"""Word document conversion using python-docx with theme support."""

import base64
import binascii
import sys
from io import BytesIO
from pathlib import Path
from typing import Optional

from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as DOCX_REL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import canonical themes - handle both package and frozen modes
try:
    from cc_shared.themes import CanonicalTheme, get_theme
except ImportError:
    try:
        from themes import CanonicalTheme, get_theme
    except ImportError:
        CanonicalTheme = None
        get_theme = None


# Maximum width for an inserted image so large screenshots do not overflow the
# printable page area.
_MAX_IMAGE_WIDTH = Inches(6)


def _hex_to_rgb(hex_color: str) -> Optional[RGBColor]:
    """Convert hex color string to RGBColor. Returns None for non-hex values."""
    hex_color = hex_color.strip()
    if not hex_color.startswith("#"):
        return None
    hex_color = hex_color.lstrip("#")
    if len(hex_color) == 3:
        hex_color = "".join(c * 2 for c in hex_color)
    if len(hex_color) != 6:
        return None
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    except ValueError:
        return None


def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    hex_clean = hex_color.lstrip("#")
    if len(hex_clean) != 6:
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_clean)
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)


def _set_paragraph_shading(para, hex_color: str) -> None:
    """Apply background shading to a paragraph."""
    hex_clean = hex_color.lstrip("#")
    if len(hex_clean) != 6:
        return
    pPr = para._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_clean)
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)


def _add_paragraph_border_bottom(para, color_hex: str, size: int = 6) -> None:
    """Add a bottom border to a paragraph (used for horizontal rules)."""
    hex_clean = color_hex.lstrip("#")
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_clean)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_paragraph_border_left(para, color_hex: str, size: int = 12) -> None:
    """Add a left border to a paragraph (used for blockquotes)."""
    hex_clean = color_hex.lstrip("#")
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_clean)
    pBdr.append(left)
    pPr.append(pBdr)


def _line_height_to_spacing(line_height_str: str) -> Optional[float]:
    """Convert CSS line-height string to a Word proportional line spacing value."""
    try:
        return float(line_height_str)
    except (ValueError, TypeError):
        return None


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a real external hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, DOCX_REL.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    new_run.append(text_el)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _decode_image_src(src: str, base_path: Optional[Path]) -> Optional[bytes]:
    """Resolve an <img> src to raw image bytes, or None if it cannot be read."""
    if not src:
        return None

    # Base64 data URI
    if src.startswith("data:"):
        try:
            header, _, b64 = src.partition(",")
            if "base64" in header:
                return base64.b64decode(b64)
        except (ValueError, binascii.Error):
            return None
        return None

    # Remote URLs are not fetched (no network in conversion path).
    if src.startswith(("http://", "https://")):
        return None

    # Local file reference
    if src.startswith("file:///"):
        img_path = Path(src[8:])
    elif base_path is not None:
        img_path = (base_path / src).resolve()
    else:
        img_path = Path(src)

    if img_path.exists():
        try:
            return img_path.read_bytes()
        except OSError:
            return None
    return None


def _add_inline_image(paragraph, img_node: Tag, base_path: Optional[Path]) -> None:
    """Insert an image inline into a paragraph; warn visibly if it cannot load."""
    src = img_node.get("src", "")
    alt = img_node.get("alt", "") or ""

    data = _decode_image_src(src, base_path)
    if data is None:
        # Visible, ASCII-only warning. Keep the alt text in the document so the
        # reader sees that something was meant to be here.
        print(f"WARNING: could not embed image {src or '(no src)'}", file=sys.stderr)
        if alt:
            paragraph.add_run(f"[image: {alt}]")
        return

    try:
        run = paragraph.add_run()
        picture = run.add_picture(BytesIO(data))
    except Exception as exc:  # python-docx raises various errors for bad images
        print(f"WARNING: could not embed image {src or '(no src)'}: {exc}", file=sys.stderr)
        if alt:
            paragraph.add_run(f"[image: {alt}]")
        return

    # Scale down oversized images while preserving aspect ratio.
    if picture.width and picture.width > _MAX_IMAGE_WIDTH:
        ratio = _MAX_IMAGE_WIDTH / picture.width
        picture.width = int(picture.width * ratio)
        picture.height = int(picture.height * ratio)


def _add_run(paragraph, text: str, theme: Optional["CanonicalTheme"],
             bold: bool, italic: bool, code: bool, strike: bool):
    """Add a styled text run to a paragraph."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if strike:
        run.font.strike = True
    if code:
        run.font.name = theme.fonts.code if theme else "Consolas"
        run.font.size = Pt(9)
    return run


def _render_inline(paragraph, node, theme: Optional["CanonicalTheme"],
                   base_path: Optional[Path],
                   bold: bool = False, italic: bool = False,
                   code: bool = False, strike: bool = False) -> None:
    """Render the inline children of an element into a Word paragraph.

    Emits text runs with the correct bold/italic/strikethrough/inline-code
    formatting, real hyperlinks for <a>, and inserted pictures for <img>.
    Block-level lists are skipped here (handled by the list processor).
    """
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                _add_run(paragraph, text, theme, bold, italic, code, strike)
            continue

        name = child.name
        if name in ("strong", "b"):
            _render_inline(paragraph, child, theme, base_path, True, italic, code, strike)
        elif name in ("em", "i"):
            _render_inline(paragraph, child, theme, base_path, bold, True, code, strike)
        elif name in ("del", "s", "strike"):
            _render_inline(paragraph, child, theme, base_path, bold, italic, code, True)
        elif name == "code":
            _render_inline(paragraph, child, theme, base_path, bold, italic, True, strike)
        elif name == "a":
            href = child.get("href", "")
            text = child.get_text()
            if href:
                _add_hyperlink(paragraph, href, text)
            elif text:
                _add_run(paragraph, text, theme, bold, italic, code, strike)
        elif name == "img":
            _add_inline_image(paragraph, child, base_path)
        elif name == "br":
            paragraph.add_run().add_break()
        elif name in ("ul", "ol"):
            # Nested lists are handled by the list processor, not inline.
            continue
        else:
            # span, p inside li, and any other inline container: recurse.
            _render_inline(paragraph, child, theme, base_path, bold, italic, code, strike)


def convert_to_word(
    html_content: str,
    output_path: Path,
    theme_name: str = "paper",
    base_path: Optional[Path] = None,
) -> None:
    """Convert HTML to a Word document with theme styling.

    Maps HTML elements to Word styles:
    - h1-h6 -> Heading 1-6
    - p -> Normal (with inline bold/italic/links/inline code/images)
    - ul/ol -> List styles (nested lists indented)
    - table -> Table with theme colors
    - pre/code -> Code style with theme code font
    - blockquote -> Quote style
    - hr -> Paragraph with bottom border
    - img -> inserted picture (resolved relative to base_path)

    Args:
        html_content: Complete HTML document string
        output_path: Path for output .docx file
        theme_name: Theme to apply for fonts and colors
        base_path: Directory used to resolve relative <img> sources
    """
    # Get theme
    theme = None
    if get_theme is not None:
        try:
            theme = get_theme(theme_name)
        except ValueError:
            pass

    # Parse HTML
    soup = BeautifulSoup(html_content, "html.parser")

    # Create document
    doc = Document()

    # Apply theme fonts and spacing to default style
    if theme:
        style = doc.styles["Normal"]
        style.font.name = theme.fonts.body
        text_color = _hex_to_rgb(theme.colors.text)
        if text_color:
            style.font.color.rgb = text_color

        line_height = _line_height_to_spacing(theme.style.body_line_height)
        if line_height:
            style.paragraph_format.line_spacing = line_height
        style.paragraph_format.space_after = Pt(6)

        for level in range(1, 7):
            style_name = f"Heading {level}"
            if style_name in [s.name for s in doc.styles]:
                h_style = doc.styles[style_name]
                h_style.font.name = theme.fonts.heading
                heading_color = _hex_to_rgb(theme.colors.heading)
                if heading_color:
                    h_style.font.color.rgb = heading_color
                h_style.paragraph_format.space_before = Pt(18 - level * 2)
                h_style.paragraph_format.space_after = Pt(6)

    # Find the main content
    body = soup.find("article") or soup.find("body") or soup

    # Process elements
    _process_element(doc, body, theme, base_path)

    # Ensure output directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Save
    doc.save(str(output_path))


def _element_has_content(element: Tag) -> bool:
    """True if the element carries visible text or an image."""
    if element.get_text(strip=True):
        return True
    return element.find("img") is not None


def _process_element(doc: Document, element, theme: Optional["CanonicalTheme"] = None,
                     base_path: Optional[Path] = None):
    """Recursively process HTML block elements."""
    if element.name is None:
        return

    for child in element.children:
        if child.name is None:
            continue

        if child.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(child.name[1])
            heading = doc.add_heading("", level=level)
            _render_inline(heading, child, theme, base_path)

            if theme and level == 1:
                primary_color = _hex_to_rgb(theme.colors.primary)
                if primary_color:
                    for run in heading.runs:
                        run.font.color.rgb = primary_color

        elif child.name == "p":
            if _element_has_content(child):
                para = doc.add_paragraph()
                _render_inline(para, child, theme, base_path)

        elif child.name == "ul":
            _process_list(doc, child, theme, base_path, ordered=False, level=0)

        elif child.name == "ol":
            _process_list(doc, child, theme, base_path, ordered=True, level=0)

        elif child.name == "blockquote":
            _process_blockquote(doc, child, theme, base_path)

        elif child.name == "pre":
            _process_code_block(doc, child, theme)

        elif child.name == "table":
            _process_table(doc, child, theme, base_path)

        elif child.name == "hr":
            _process_hr(doc, theme)

        elif child.name in ["div", "article", "section", "main"]:
            _process_element(doc, child, theme, base_path)


def _process_list(doc: Document, list_element, theme: Optional["CanonicalTheme"],
                  base_path: Optional[Path], ordered: bool = False, level: int = 0):
    """Process a ul or ol list, recursing for nested lists with indentation."""
    for li in list_element.find_all("li", recursive=False):
        style = "List Number" if ordered else "List Bullet"
        try:
            para = doc.add_paragraph(style=style)
        except KeyError:
            para = doc.add_paragraph()

        # Indent nested levels.
        if level > 0:
            para.paragraph_format.left_indent = Inches(0.25 * (level + 1))

        # Render the item's own inline content (nested lists are skipped here).
        _render_inline(para, li, theme, base_path)

        # Recurse into any nested lists.
        for sub in li.find_all(["ul", "ol"], recursive=False):
            _process_list(doc, sub, theme, base_path,
                          ordered=(sub.name == "ol"), level=level + 1)


def _process_blockquote(doc: Document, element, theme: Optional["CanonicalTheme"] = None,
                        base_path: Optional[Path] = None):
    """Process blockquote with left border and optional italic styling."""
    if not _element_has_content(element):
        return

    para = doc.add_paragraph()
    _render_inline(para, element, theme, base_path)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    if theme:
        border_color = theme.colors.blockquote_border
        if border_color.startswith("#"):
            _add_paragraph_border_left(para, border_color)

        bq_color = _hex_to_rgb(theme.colors.blockquote_text)
        if bq_color:
            for run in para.runs:
                run.font.color.rgb = bq_color

        if theme.name in ("boardroom", "thesis"):
            for run in para.runs:
                run.font.italic = True


def _process_code_block(doc: Document, element, theme: Optional["CanonicalTheme"] = None):
    """Process code block with monospace font and background shading."""
    code_text = element.get_text()
    para = doc.add_paragraph()
    run = para.add_run(code_text)

    if theme:
        run.font.name = theme.fonts.code
    else:
        run.font.name = "Consolas"
    run.font.size = Pt(9)

    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.right_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    if theme:
        code_bg = theme.colors.code_bg
        if code_bg.startswith("#"):
            _set_paragraph_shading(para, code_bg)

        code_color = _hex_to_rgb(theme.colors.code_text)
        if code_color:
            run.font.color.rgb = code_color


def _process_hr(doc: Document, theme: Optional["CanonicalTheme"] = None):
    """Process horizontal rule as a paragraph with bottom border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)

    if theme:
        border_color = theme.colors.border
        if theme.name == "boardroom":
            border_color = theme.colors.accent
        if border_color.startswith("#"):
            _add_paragraph_border_bottom(para, border_color, size=6)
    else:
        _add_paragraph_border_bottom(para, "CCCCCC", size=6)


def _process_table(doc: Document, table_element, theme: Optional["CanonicalTheme"] = None,
                   base_path: Optional[Path] = None):
    """Process HTML table with theme colors and alternating rows."""
    rows = table_element.find_all("tr")
    if not rows:
        return

    first_row = rows[0]
    cols = first_row.find_all(["th", "td"])
    num_cols = len(cols)

    if num_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for col_idx, cell in enumerate(cells):
            if col_idx >= num_cols:
                continue

            doc_cell = table.rows[row_idx].cells[col_idx]
            # Render inline content into the cell's existing first paragraph.
            para = doc_cell.paragraphs[0]
            _render_inline(para, cell, theme, base_path)

            if cell.name == "th":
                for paragraph in doc_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        if theme:
                            header_text_color = _hex_to_rgb(theme.colors.table_header_text)
                            if header_text_color:
                                run.font.color.rgb = header_text_color
                            run.font.name = theme.fonts.heading

                if theme:
                    header_bg = theme.colors.table_header_bg
                    if header_bg.startswith("#"):
                        _set_cell_shading(doc_cell, header_bg)

            elif theme and cell.name == "td" and row_idx % 2 == 0 and row_idx > 0:
                alt_bg = theme.colors.alt_row_bg
                if alt_bg.startswith("#"):
                    _set_cell_shading(doc_cell, alt_bg)
