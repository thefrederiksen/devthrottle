"""Tests for cc-word Word converter."""

import sys
from pathlib import Path

# Add src and cc_shared to path for testing
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))
sys.path.insert(0, str(Path(__file__).parent.parent.parent / "cc_shared"))

import base64
import struct
import zlib

from word_converter import _hex_to_rgb, convert_to_word
from docx import Document
from docx.shared import RGBColor


def _make_png() -> bytes:
    """Build a valid 1x1 grayscale PNG for inline-image tests."""
    def _chunk(typ: bytes, data: bytes) -> bytes:
        body = typ + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 0, 0, 0, 0)
    raw_scanline = b"\x00\x00"  # filter byte + one gray pixel
    idat = zlib.compress(raw_scanline)
    return signature + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")


# A minimal valid 1x1 PNG, used to test inline image insertion.
_TINY_PNG = _make_png()


def _convert(html, tmp_path, **kwargs):
    out = tmp_path / "out.docx"
    convert_to_word(html, out, **kwargs)
    return Document(str(out))


class TestHexToRgb:
    def test_valid_6_digit_hex(self):
        result = _hex_to_rgb("#1A365D")
        assert result == RGBColor(0x1A, 0x36, 0x5D)

    def test_valid_3_digit_hex(self):
        result = _hex_to_rgb("#FFF")
        assert result == RGBColor(0xFF, 0xFF, 0xFF)

    def test_lowercase_hex(self):
        result = _hex_to_rgb("#abc123")
        assert result == RGBColor(0xAB, 0xC1, 0x23)

    def test_no_hash_returns_none(self):
        result = _hex_to_rgb("1A365D")
        assert result is None

    def test_rgba_returns_none(self):
        result = _hex_to_rgb("rgba(168, 85, 247, 0.1)")
        assert result is None

    def test_empty_returns_none(self):
        result = _hex_to_rgb("")
        assert result is None


class TestConvertToWord:
    def test_basic_conversion(self, tmp_path):
        html = "<article><h1>Title</h1><p>Body text</p></article>"
        output = tmp_path / "test.docx"
        convert_to_word(html, output, theme_name="paper")
        assert output.exists()
        assert output.stat().st_size > 0

    def test_with_table(self, tmp_path):
        html = """<article>
        <table><tr><th>Name</th><th>Value</th></tr>
        <tr><td>A</td><td>1</td></tr></table>
        </article>"""
        output = tmp_path / "table.docx"
        convert_to_word(html, output, theme_name="boardroom")
        assert output.exists()

    def test_with_code_block(self, tmp_path):
        html = "<article><pre><code>print('hello')</code></pre></article>"
        output = tmp_path / "code.docx"
        convert_to_word(html, output, theme_name="terminal")
        assert output.exists()

    def test_with_list(self, tmp_path):
        html = "<article><ul><li>Item 1</li><li>Item 2</li></ul></article>"
        output = tmp_path / "list.docx"
        convert_to_word(html, output, theme_name="paper")
        assert output.exists()

    def test_invalid_theme_still_works(self, tmp_path):
        html = "<article><p>Text</p></article>"
        output = tmp_path / "fallback.docx"
        convert_to_word(html, output, theme_name="nonexistent")
        assert output.exists()

    def test_creates_output_directory(self, tmp_path):
        html = "<article><p>Text</p></article>"
        output = tmp_path / "subdir" / "test.docx"
        convert_to_word(html, output, theme_name="paper")
        assert output.exists()


class TestInlineFormatting:
    def test_bold_and_italic_runs(self, tmp_path):
        html = "<article><p><strong>boldword</strong> and <em>italword</em></p></article>"
        doc = _convert(html, tmp_path, theme_name="paper")
        runs = [r for p in doc.paragraphs for r in p.runs]
        bold = [r for r in runs if r.text == "boldword"]
        ital = [r for r in runs if r.text == "italword"]
        assert bold and bold[0].bold is True
        assert ital and ital[0].italic is True

    def test_inline_code_uses_monospace(self, tmp_path):
        html = "<article><p>see <code>x = 1</code> here</p></article>"
        doc = _convert(html, tmp_path, theme_name="terminal")
        code_runs = [r for p in doc.paragraphs for r in p.runs if r.text == "x = 1"]
        assert code_runs
        # Monospace code font applied (not the default body font).
        assert code_runs[0].font.name is not None

    def test_hyperlink_relationship_created(self, tmp_path):
        html = '<article><p><a href="https://example.com/page">link</a></p></article>'
        doc = _convert(html, tmp_path, theme_name="paper")
        targets = [rel.target_ref for rel in doc.part.rels.values()]
        assert any("example.com/page" in t for t in targets)

    def test_inline_image_inserted(self, tmp_path):
        b64 = base64.b64encode(_TINY_PNG).decode("ascii")
        html = f'<article><p><img src="data:image/png;base64,{b64}" alt="x"></p></article>'
        doc = _convert(html, tmp_path, theme_name="paper")
        assert len(doc.inline_shapes) == 1

    def test_local_image_inserted_via_base_path(self, tmp_path):
        img = tmp_path / "pic.png"
        img.write_bytes(_TINY_PNG)
        html = '<article><p><img src="pic.png" alt="x"></p></article>'
        out = tmp_path / "out.docx"
        convert_to_word(html, out, theme_name="paper", base_path=tmp_path)
        doc = Document(str(out))
        assert len(doc.inline_shapes) == 1

    def test_missing_image_warns_not_silent(self, tmp_path, capsys):
        html = '<article><p><img src="gone.png" alt="caption"></p></article>'
        out = tmp_path / "out.docx"
        convert_to_word(html, out, theme_name="paper", base_path=tmp_path)
        captured = capsys.readouterr()
        assert "WARNING" in captured.err
        assert captured.err.isascii()


class TestNestedLists:
    def test_nested_list_is_indented(self, tmp_path):
        html = (
            "<article><ul><li>Top<ul><li>Nested</li></ul></li></ul></article>"
        )
        doc = _convert(html, tmp_path, theme_name="paper")
        texts = {p.text: p for p in doc.paragraphs if p.text}
        assert "Top" in texts
        assert "Nested" in texts
        # The nested item carries a left indent; the top-level item does not.
        nested = texts["Nested"]
        assert nested.paragraph_format.left_indent is not None
        assert nested.paragraph_format.left_indent > 0
